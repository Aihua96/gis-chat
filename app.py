#!/usr/bin/env python3
"""零依赖 Web MVP：工程测量矢量成果到技术总结报告。"""
from http.server import ThreadingHTTPServer, SimpleHTTPRequestHandler
from pathlib import Path
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from datetime import date
import base64, io, json, math, os, re, shutil, struct, uuid, zipfile

ROOT = Path(__file__).resolve().parent
DATA = ROOT / "data"
OUTPUT = ROOT / "output"
OUTPUT.mkdir(exist_ok=True)
PYDOCX_OK = True
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    PYDOCX_OK = False

PROJECT = {
    "name": "新城产业园西侧配套道路竣工测量项目",
    "type": "工程测量（道路竣工测量）",
    "location": "四川省成都市新城产业园",
    "period": "2026年7月15日—2026年7月22日",
    "coordinate_system": "CGCS2000 / 3度分带高斯-克吕格投影（37带）",
    "unit": "四川地质十一队（演示）",
    "instrument": "GNSS 接收机、全站仪、无人机航测系统",
}

def dbf_records(raw):
    """Read the dBase attribute table bundled with an ESRI Shapefile."""
    if not raw: return []
    record_count, header_len, record_len = struct.unpack('<xxxxIHH20x', raw[:32])
    fields, offset = [], 32
    while offset < header_len and raw[offset] != 0x0D:
        name = raw[offset:offset+11].split(b'\0', 1)[0].decode('latin1').strip()
        fields.append((name, chr(raw[offset+11]), raw[offset+16])); offset += 32
    records = []
    for i in range(record_count):
        row = raw[header_len+i*record_len:header_len+(i+1)*record_len]
        if not row or row[0:1] == b'*': continue
        pos, item = 1, {}
        for name, typ, width in fields:
            value = row[pos:pos+width].decode('gb18030', errors='replace').strip(); pos += width
            if typ in ('N','F') and value:
                try: value = float(value) if '.' in value else int(value)
                except ValueError: pass
            item[name.lower()] = value
        records.append(item)
    return records

def parse_shp(raw, attributes, layer_name):
    """Parse Point / PolyLine / Polygon records from a .shp stream (MVP subset)."""
    if len(raw) < 100: raise ValueError(f'{layer_name}.shp 文件无效或内容不完整')
    features, offset, index = [], 100, 0
    while offset + 8 <= len(raw):
        _, words = struct.unpack('>2i', raw[offset:offset+8]); offset += 8
        body = raw[offset:offset+words*2]; offset += words*2
        if len(body) < 4: continue
        kind = struct.unpack('<i', body[:4])[0]; prop = dict(attributes[index]) if index < len(attributes) else {}; index += 1
        prop['layer'] = prop.get('layer') or layer_name
        prop['id'] = str(prop.get('id') or prop.get('name') or f'{layer_name}-{index:03d}')
        if kind == 0: continue
        if kind == 1:
            x, y = struct.unpack('<2d', body[4:20]); geometry = {'type':'Point','coordinates':[x,y]}
        elif kind in (3,5):
            _, _, _, _, parts, points = struct.unpack('<4d2i', body[4:44])
            starts = list(struct.unpack('<'+'i'*parts, body[44:44+4*parts]))
            start = 44 + 4*parts
            xy = [list(struct.unpack('<2d', body[start+i*16:start+(i+1)*16])) for i in range(points)]
            chunks = [xy[s:(starts[p+1] if p+1 < parts else len(xy))] for p,s in enumerate(starts)]
            geometry = {'type':'LineString','coordinates':chunks[0]} if kind == 3 and len(chunks)==1 else ({'type':'MultiLineString','coordinates':chunks} if kind == 3 else {'type':'Polygon','coordinates':chunks})
        else:
            prop['parse_warning'] = f'未支持 ShapeType {kind}'
            continue
        features.append({'type':'Feature','properties':prop,'geometry':geometry})
    return features

def shpzip_to_geojson(encoded):
    try: archive = zipfile.ZipFile(io.BytesIO(base64.b64decode(encoded)))
    except Exception as exc: raise ValueError('无法读取 ZIP，请确认压缩包包含 Shapefile 文件') from exc
    names = archive.namelist(); shp_names = [n for n in names if n.lower().endswith('.shp')]
    if not shp_names: raise ValueError('压缩包未找到 .shp 文件')
    features = []
    for shp_name in shp_names:
        base = shp_name.rsplit('.',1)[0]; dbf_name = next((n for n in names if n.rsplit('.',1)[0].lower()==base.lower() and n.lower().endswith('.dbf')), None)
        attrs = dbf_records(archive.read(dbf_name)) if dbf_name else []
        features.extend(parse_shp(archive.read(shp_name), attrs, Path(base).name))
    return {'type':'FeatureCollection','name':'SHP 导入成果','features':features,
            'crs':{'type':'name','properties':{'name':PROJECT['coordinate_system']}}}

def extract_material(filename, encoded):
    raw = base64.b64decode(encoded)
    suffix = Path(filename).suffix.lower()
    if suffix in ('.txt', '.md', '.csv'):
        return raw.decode('utf-8', errors='replace')[:30000]
    if suffix == '.docx' and PYDOCX_OK:
        doc = Document(io.BytesIO(raw)); chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            chunks.extend(' | '.join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return '\n'.join(chunks)[:30000]
    raise ValueError('资料仅支持 TXT、MD、CSV、DOCX；PDF 将在生产版接入专业解析服务')

def fallback_sections(data, description, material):
    a = analyse(data)
    project_note = description.strip() or '用户未补充项目描述，以下内容依据导入矢量成果生成。'
    material_note = '已参考上传项目资料。' if material.strip() else '未上传文本资料，建议补充任务书、设计书或作业记录。'
    return {
        '项目概况': f'{project_note}。本项目处理 {a["feature_count"]} 个空间要素，空间参考为 {a["crs"]}。{material_note}',
        '技术路线': '外业成果经数据导入后，由 GIS 引擎完成几何类型识别、图层统计、面积长度计算及基础规则检查；系统将确定性计算结果作为报告数据来源。',
        '质量评价': ('系统检出 '+str(len(a['warnings']))+' 项需人工复核事项：'+'；'.join(a['warnings'])+'。' if a['warnings'] else '基础属性完整性与支持的几何类型检查未发现异常。')+' 本评价为 AI 辅助初稿，最终结论须由项目负责人审核确认。',
        '结论与建议': '建议在归档前完成异常项复核，并核验项目资料、图件编号、精度指标及最终签章信息。'
    }

def llm_sections(data, description, material, config):
    """Generate narrative-only report sections using an OpenAI-compatible endpoint."""
    if not config or not config.get('base_url') or not config.get('model'):
        return fallback_sections(data, description, material), '演示模式（未配置大模型）'
    a = analyse(data)
    evidence = json.dumps({'statistics':a, 'project':PROJECT}, ensure_ascii=False)
    prompt = f'''你是测绘工程技术报告助手。根据项目描述、项目资料和 GIS 引擎证据，生成中文报告的分析性章节。严格规则：
1) 只能引用 GIS 证据中的数值；不得自行计算或杜撰面积、长度、精度、坐标、日期或规范条款。
2) 对缺失信息明确写“待人工补充/核验”，不要猜测。
3) 不要输出 Markdown、标题编号、免责声明或数据表；只返回 JSON 对象，键必须为：项目概况、技术路线、质量评价、结论与建议。
4) 每个值为 80—220 字的正式中文段落。

用户描述：{description or '未提供'}
项目资料摘录：{material[:12000] or '未提供'}
GIS 引擎证据（可信数值来源）：{evidence}'''
    base = config['base_url'].rstrip('/')
    endpoint = base if base.endswith('/chat/completions') else base + ('/chat/completions' if base.endswith('/v1') else '/v1/chat/completions')
    headers={'Content-Type':'application/json'}
    if config.get('api_key'): headers['Authorization']='Bearer '+config['api_key']
    body=json.dumps({'model':config['model'],'messages':[{'role':'system','content':'你只能基于给定证据编写测绘技术报告。'},{'role':'user','content':prompt}],'temperature':0.2,'response_format':{'type':'json_object'}}).encode()
    try:
        with urlopen(Request(endpoint, data=body, headers=headers, method='POST'), timeout=90) as response:
            result=json.loads(response.read().decode())
        content=result['choices'][0]['message']['content']
        sections=json.loads(re.sub(r'^```json\s*|\s*```$', '', content.strip()))
        required=['项目概况','技术路线','质量评价','结论与建议']
        if not all(isinstance(sections.get(k), str) for k in required): raise ValueError('模型未按约定返回完整章节')
        return {k:sections[k].strip() for k in required}, '大模型生成：'+config['model']
    except Exception as exc:
        raise RuntimeError('大模型调用失败：'+str(exc)) from exc

def ring_area(coords):
    if len(coords) < 3: return 0.0
    return abs(sum(coords[i][0] * coords[(i+1)%len(coords)][1] - coords[(i+1)%len(coords)][0] * coords[i][1] for i in range(len(coords))) / 2)

def line_length(coords):
    return sum(math.hypot(coords[i+1][0]-coords[i][0], coords[i+1][1]-coords[i][1]) for i in range(len(coords)-1))

def bbox(geometry):
    points = []
    def walk(v):
        if isinstance(v, (list, tuple)) and v and isinstance(v[0], (int, float)): points.append(v)
        elif isinstance(v, (list, tuple)):
            for x in v: walk(x)
    walk(geometry.get("coordinates", []))
    if not points: return None
    xs, ys = [p[0] for p in points], [p[1] for p in points]
    return min(xs), min(ys), max(xs), max(ys)

def analyse(data):
    features = data.get("features", [])
    layers, warnings, ids, polygons = {}, [], set(), []
    total_area = total_length = 0.0
    for index, f in enumerate(features, 1):
        prop, geom = f.get("properties") or {}, f.get("geometry") or {}
        layer = prop.get("layer") or "未分类图层"
        record = layers.setdefault(layer, {"count": 0, "area": 0.0, "length": 0.0})
        record["count"] += 1
        feature_id = prop.get("id")
        if not feature_id: warnings.append(f"第 {index} 个要素缺少 id 属性")
        elif feature_id in ids: warnings.append(f"要素编号重复：{feature_id}")
        else: ids.add(feature_id)
        typ, c = geom.get("type"), geom.get("coordinates", [])
        if typ == "Polygon":
            area = ring_area(c[0]) - sum(ring_area(hole) for hole in c[1:])
            record["area"] += area; total_area += area; polygons.append((feature_id or str(index), bbox(geom)))
        elif typ == "LineString":
            length = line_length(c); record["length"] += length; total_length += length
        elif typ == "MultiLineString":
            length = sum(line_length(part) for part in c); record["length"] += length; total_length += length
        elif typ not in ("Point",): warnings.append(f"第 {index} 个要素为当前 MVP 未支持的几何类型：{typ}")
        if prop.get("quality") == "待复核": warnings.append(f"{feature_id or '未编号要素'} 标记为待复核")
    # 演示级 bbox 重叠检查：仅做预警，不能替代严格拓扑验证。
    for i in range(len(polygons)):
        for j in range(i+1, len(polygons)):
            a, b = polygons[i][1], polygons[j][1]
            if a and b and max(a[0],b[0]) < min(a[2],b[2]) and max(a[1],b[1]) < min(a[3],b[3]):
                warnings.append(f"面要素包络可能重叠：{polygons[i][0]} / {polygons[j][0]}（需人工复核）")
    return {"feature_count": len(features), "layers": layers, "total_area": total_area,
            "total_length": total_length, "warnings": warnings,
            "crs": (data.get("crs") or {}).get("properties", {}).get("name") or PROJECT["coordinate_system"]}

def svg_preview(data):
    all_bbox = bbox({"coordinates": [f.get("geometry",{}).get("coordinates",[]) for f in data.get("features",[])]})
    if not all_bbox: return '<svg xmlns="http://www.w3.org/2000/svg" width="800" height="430"></svg>'
    x0,y0,x1,y1 = all_bbox; dx,dy=max(x1-x0,1),max(y1-y0,1)
    def p(pt): return ((pt[0]-x0)/dx*700+50, 380-(pt[1]-y0)/dy*320)
    shape = ['<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 800 430"><rect width="800" height="430" fill="#f6f8fb"/><text x="50" y="32" font-family="Arial" font-size="16" fill="#0f2747">测区成果示意图（自动生成）</text>']
    colors = {"项目范围":"#2e73b8", "道路中心线":"#e76f51", "控制点":"#264653", "绿化边界":"#65a30d"}
    for f in data.get("features",[]):
        g, prop = f.get("geometry",{}), f.get("properties",{})
        color=colors.get(prop.get("layer"), "#64748b"); typ=g.get("type"); c=g.get("coordinates",[])
        if typ=="Polygon":
            pts=' '.join(f'{a:.1f},{b:.1f}' for a,b in map(p,c[0])); shape.append(f'<polygon points="{pts}" fill="{color}" fill-opacity=".12" stroke="{color}" stroke-width="2"/>')
        if typ=="LineString":
            pts=' '.join(f'{a:.1f},{b:.1f}' for a,b in map(p,c)); shape.append(f'<polyline points="{pts}" fill="none" stroke="{color}" stroke-width="5" stroke-linecap="round"/>')
        if typ=="MultiLineString":
            for part in c:
                pts=' '.join(f'{a:.1f},{b:.1f}' for a,b in map(p,part)); shape.append(f'<polyline points="{pts}" fill="none" stroke="{color}" stroke-width="5" stroke-linecap="round"/>')
        if typ=="Point":
            a,b=p(c); shape.append(f'<circle cx="{a:.1f}" cy="{b:.1f}" r="6" fill="{color}"/><text x="{a+8:.1f}" y="{b-8:.1f}" font-family="Arial" font-size="11" fill="#243b53">{prop.get("id","")}</text>')
    shape.append('<text x="50" y="410" font-family="Arial" font-size="11" fill="#52616b">蓝：测区边界　红：道路中心线　绿：绿化边界　深色点：控制点</text></svg>')
    return ''.join(shape)

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def build_report(data, output_path, description='', material='', template='engineering_summary', config=None):
    if not PYDOCX_OK: raise RuntimeError('未找到 python-docx 运行依赖')
    a = analyse(data); sections, ai_status = llm_sections(data, description, material, config); doc = Document(); section = doc.sections[0]
    for attr in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(section, attr, Cm(2.54))
    # Heiti SC is available in the macOS runtime used for visual QA; clients can
    # replace it with their approved CJK font through the production template.
    normal = doc.styles['Normal']; normal.font.name='Heiti SC'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Heiti SC'); normal.font.size=Pt(10.5)
    normal.paragraph_format.line_spacing=1.25; normal.paragraph_format.space_after=Pt(5)
    for style_name, size, color in [('Heading 1',16,'1F4D78'),('Heading 2',13,'2E74B5')]:
        s=doc.styles[style_name]; s.font.name='Heiti SC'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Heiti SC'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(80)
    title = '工程测量技术设计书' if template == 'engineering_design' else '工程测量技术总结报告'
    r=p.add_run(title); r.bold=True; r.font.name='Heiti SC'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Heiti SC'); r.font.size=Pt(24); r.font.color.rgb=RGBColor(15,55,95)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(22); p.add_run(PROJECT['name']).bold=True
    for text in [PROJECT['unit'], '报告生成日期：'+str(date.today()), '（MVP 演示样本，非正式测绘成果）']:
        p=doc.add_paragraph(text); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading('一、项目概况', 1)
    doc.add_paragraph(sections['项目概况'])
    doc.add_heading('二、数据来源与处理口径', 1)
    doc.add_paragraph('本报告的数据统计由系统确定性 GIS 计算模块生成。面积按投影平面坐标的多边形鞋带公式计算，线长度按相邻顶点欧氏距离累计计算；文本由受控模板生成。')
    doc.add_heading('三、技术路线', 1)
    doc.add_paragraph(sections['技术路线'])
    doc.add_heading('四、成果统计', 1)
    table=doc.add_table(rows=1, cols=4); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    heads=['图层','要素数','面积（㎡）','长度（m）']
    for cell, text in zip(table.rows[0].cells, heads):
        cell.text=text; set_cell_shading(cell,'E8EEF5'); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for name, v in a['layers'].items():
        cells=table.add_row().cells
        for c,text in zip(cells,[name,str(v['count']),f"{v['area']:.2f}" if v['area'] else '—',f"{v['length']:.2f}" if v['length'] else '—']): c.text=text
    doc.add_paragraph(f"汇总：共处理 {a['feature_count']} 个要素；面状成果面积 {a['total_area']:.2f} ㎡（{a['total_area']/10000:.4f} ha）；线状成果长度 {a['total_length']:.2f} m。", style='Normal')
    doc.add_heading('五、质量检查与评价', 1)
    if a['warnings']:
        doc.add_paragraph('系统发现以下需复核项：')
        for item in a['warnings']: doc.add_paragraph(item, style='List Bullet')
    else: doc.add_paragraph('基础属性完整性与几何类型检查未发现异常。')
    doc.add_paragraph(sections['质量评价'])
    doc.add_heading('六、结论与建议', 1)
    doc.add_paragraph(sections['结论与建议'])
    doc.add_heading('附录 A：数据血缘记录', 1)
    doc.add_paragraph(f"源文件：用户上传 GIS 矢量成果；空间参考：{a['crs']}；生成时间：{date.today()}；处理版本：MVP 0.3；文本生成：{ai_status}。")
    doc.save(output_path)

class Handler(SimpleHTTPRequestHandler):
    def send_json(self, payload, status=200):
        content=json.dumps(payload, ensure_ascii=False).encode(); self.send_response(status); self.send_header('Content-Type','application/json; charset=utf-8'); self.send_header('Content-Length',str(len(content))); self.end_headers(); self.wfile.write(content)
    def do_GET(self):
        path=urlparse(self.path).path
        demo=json.loads((DATA/'demo_project.geojson').read_text())
        if path=='/api/project': return self.send_json({**PROJECT,'analysis':analyse(demo)})
        if path=='/api/preview.svg':
            content=svg_preview(demo).encode(); self.send_response(200); self.send_header('Content-Type','image/svg+xml'); self.send_header('Content-Length',str(len(content))); self.end_headers(); return self.wfile.write(content)
        if path=='/api/demo.geojson':
            content=(DATA/'demo_project.geojson').read_bytes(); self.send_response(200); self.send_header('Content-Type','application/geo+json'); self.send_header('Content-Length',str(len(content))); self.end_headers(); return self.wfile.write(content)
        if path=='/api/demo-shp.zip':
            content=(DATA/'demo_engineering_survey_shp.zip').read_bytes(); self.send_response(200); self.send_header('Content-Type','application/zip'); self.send_header('Content-Disposition','attachment; filename=demo_engineering_survey_shp.zip'); self.send_header('Content-Length',str(len(content))); self.end_headers(); return self.wfile.write(content)
        if path.startswith('/output/'):
            return super().do_GET()
        if path in ('/','/index.html'):
            self.path='/web/index.html'; return super().do_GET()
        return super().do_GET()
    def do_POST(self):
        path=urlparse(self.path).path
        length=int(self.headers.get('Content-Length','0')); raw=self.rfile.read(length)
        try: payload=json.loads(raw)
        except Exception: return self.send_json({'error':'请求数据无法解析'},400)
        if path == '/api/extract-material':
            try: return self.send_json({'text':extract_material(payload['filename'], payload['content'])})
            except Exception as e: return self.send_json({'error':str(e)},400)
        try:
            source = payload.get('data', payload)
            data = shpzip_to_geojson(source['content']) if source.get('format') == 'shpzip' else source
        except Exception as e: return self.send_json({'error':str(e)},400)
        if path=='/api/convert': return self.send_json(data)
        if path=='/api/analyse': return self.send_json({'analysis':analyse(data),'preview_svg':svg_preview(data)})
        if path=='/api/report':
            try:
                name='工程测量技术报告_'+uuid.uuid4().hex[:8]+'.docx'
                build_report(data, OUTPUT/name, payload.get('description',''), payload.get('material',''), payload.get('template','engineering_summary'), payload.get('llm'))
                return self.send_json({'url':'/output/'+name, 'mode':'llm' if payload.get('llm',{}).get('base_url') else 'demo'})
            except Exception as e: return self.send_json({'error':str(e)},500)
        return self.send_json({'error':'未找到接口'},404)

if __name__=='__main__':
    os.chdir(ROOT); print('MVP started: http://127.0.0.1:8080'); ThreadingHTTPServer(('127.0.0.1',8080), Handler).serve_forever()
