"""项目资料的文本提取：把用户上传的文件转成模型可用的纯文本。"""
import base64, io
from pathlib import Path

from ..config import MATERIAL_LIMIT
from ..domain import DependencyMissing, ValidationError

try:
    from docx import Document
except ImportError:
    Document = None

TEXT_SUFFIXES = ('.txt', '.md', '.csv')
SUPPORTED_SUFFIXES = TEXT_SUFFIXES + ('.docx',)

def extract_text(filename: str, encoded: str) -> str:
    suffix = Path(filename or '').suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValidationError('资料仅支持 TXT、MD、CSV、DOCX；PDF 将在生产版接入专业解析服务')
    try:
        raw = base64.b64decode(encoded or '')
    except Exception as exc:
        raise ValidationError('资料内容无法解码') from exc
    if suffix in TEXT_SUFFIXES:
        return raw.decode('utf-8', errors='replace')[:MATERIAL_LIMIT]
    if Document is None:
        raise DependencyMissing('解析 DOCX 需要 python-docx，请先执行 pip install -r requirements.txt')
    try:
        document = Document(io.BytesIO(raw))
    except Exception as exc:
        raise ValidationError('DOCX 文件无法解析，请确认文件完整') from exc
    chunks = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        chunks.extend(' | '.join(cell.text.strip() for cell in row.cells) for row in table.rows)
    return '\n'.join(chunks)[:MATERIAL_LIMIT]
