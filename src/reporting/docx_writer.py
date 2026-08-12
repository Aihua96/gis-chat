"""Word 报告排版。输入是已经确定的数据（项目、证据、叙述章节），本模块不做任何计算。"""
from datetime import date

from ..config import REPORT_FONT
from ..domain import DependencyMissing, Evidence, Narrative, Project
from .templates import ReportTemplate

DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    DOCX_AVAILABLE = False

HEADER_FILL = 'E8EEF5'

def _apply_font(item, size=None, color=None):
    item.font.name = REPORT_FONT
    item._element.rPr.rFonts.set(qn('w:eastAsia'), REPORT_FONT)
    if size: item.font.size = Pt(size)
    if color: item.font.color.rgb = RGBColor.from_string(color)

def _shade(cell, color):
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def _table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text; _shade(cell, HEADER_FILL); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table

def _setup_styles(doc):
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(doc.sections[0], attr, Cm(2.54))
    normal = doc.styles['Normal']
    _apply_font(normal, 10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [('Heading 1', 16, '1F4D78'), ('Heading 2', 13, '2E74B5')]:
        _apply_font(doc.styles[name], size, color)

def _cover(doc, project: Project, template: ReportTemplate):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(80)
    run = paragraph.add_run(template.title)
    run.bold = True
    _apply_font(run, 24)
    run.font.color.rgb = RGBColor(15, 55, 95)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(22)
    paragraph.add_run(project.name).bold = True
    for text in [project.unit or '编制单位待补充', '报告生成日期：' + str(date.today()), '（AI 辅助初稿，需人工审核后方可作为正式成果）']:
        line = doc.add_paragraph(text)
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def _project_table(doc, project: Project, evidence: Evidence):
    rows = [('项目名称', project.name), ('项目类型', project.type), ('测区位置', project.location),
            ('作业周期', project.period), ('承担单位', project.unit), ('主要仪器', project.instrument),
            ('坐标系统', evidence.crs)]
    table = _table(doc, ['项目要素', '内容'])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value or '待人工补充'

def _statistics_table(doc, evidence: Evidence):
    table = _table(doc, ['图层', '要素数', '面积（㎡）', '长度（m）'])
    for name, stat in evidence.layers.items():
        cells = table.add_row().cells
        values = [name, str(stat.count),
                  f'{stat.area:.2f}' if stat.area else '—',
                  f'{stat.length:.2f}' if stat.length else '—']
        for cell, text in zip(cells, values):
            cell.text = text

def render(project: Project, evidence: Evidence, narrative: Narrative,
           template: ReportTemplate, output_path) -> None:
    if not DOCX_AVAILABLE:
        raise DependencyMissing('生成 Word 报告需要 python-docx，请先执行 pip install -r requirements.txt')
    doc = Document()
    _setup_styles(doc)
    _cover(doc, project, template)
    doc.add_heading('一、项目概况', 1)
    _project_table(doc, project, evidence)
    doc.add_paragraph(narrative.sections['项目概况'])
    doc.add_heading('二、数据来源与处理口径', 1)
    doc.add_paragraph('本报告的数据统计由系统确定性 GIS 计算模块生成。面积按投影平面坐标的多边形鞋带公式计算，'
                      '线长度按相邻顶点欧氏距离累计计算。分析性文字由 AI 依据上述数值组织，不参与任何数值计算。')
    doc.add_heading('三、技术路线', 1)
    doc.add_paragraph(narrative.sections['技术路线'])
    doc.add_heading('四、成果统计', 1)
    _statistics_table(doc, evidence)
    doc.add_paragraph(f'汇总：共处理 {evidence.feature_count} 个要素；'
                      f'面状成果面积 {evidence.total_area:.2f} ㎡（{evidence.total_area/10000:.4f} ha）；'
                      f'线状成果长度 {evidence.total_length:.2f} m。')
    doc.add_heading('五、质量检查与评价', 1)
    if evidence.warnings:
        doc.add_paragraph('系统发现以下需复核项：')
        for warning in evidence.warnings:
            doc.add_paragraph(warning, style='List Bullet')
    else:
        doc.add_paragraph('基础属性完整性与几何类型检查未发现异常。')
    doc.add_paragraph(narrative.sections['质量评价'])
    doc.add_heading('六、结论与建议', 1)
    doc.add_paragraph(narrative.sections['结论与建议'])
    doc.add_heading('附录 A：数据血缘记录', 1)
    source = project.dataset.source if project.dataset else '未记录'
    doc.add_paragraph(f'源文件：{source}；空间参考：{evidence.crs}；报告模板：{template.name}；'
                      f'生成时间：{date.today()}；文本生成方式：{narrative.source}。')
    doc.save(output_path)
