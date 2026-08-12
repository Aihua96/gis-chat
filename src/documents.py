"""Word 文档读写：提取上传资料的文本，以及排版最终的技术报告。

python-docx 为可选依赖：缺失时应用仍可启动并完成 GIS 统计，只有报告导出不可用。
"""
import base64, io
from datetime import date
from pathlib import Path

from .config import PROJECT, REPORT_FONT

PYDOCX_OK = True
try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    PYDOCX_OK = False

TEXT_SUFFIXES = ('.txt', '.md', '.csv')
MATERIAL_LIMIT = 30000

def extract_material(filename, encoded):
    """把上传的项目资料转成纯文本，供大模型作为参考上下文。"""
    raw = base64.b64decode(encoded)
    suffix = Path(filename).suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return raw.decode('utf-8', errors='replace')[:MATERIAL_LIMIT]
    if suffix == '.docx' and PYDOCX_OK:
        doc = Document(io.BytesIO(raw))
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            chunks.extend(' | '.join(cell.text.strip() for cell in row.cells) for row in table.rows)
        return '\n'.join(chunks)[:MATERIAL_LIMIT]
    raise ValueError('资料仅支持 TXT、MD、CSV、DOCX；PDF 将在生产版接入专业解析服务')

def _use_report_font(item, size=None, color=None):
    item.font.name = REPORT_FONT
    item._element.rPr.rFonts.set(qn('w:eastAsia'), REPORT_FONT)
    if size: item.font.size = Pt(size)
    if color: item.font.color.rgb = RGBColor.from_string(color)

def _shade(cell, color):
    shading = OxmlElement('w:shd'); shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def _cover(doc, template):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(80)
    title = '工程测量技术设计书' if template == 'engineering_design' else '工程测量技术总结报告'
    run = p.add_run(title); run.bold = True; _use_report_font(run, 24); run.font.color.rgb = RGBColor(15, 55, 95)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(22)
    p.add_run(PROJECT['name']).bold = True
    for text in [PROJECT['unit'], '报告生成日期：'+str(date.today()), '（MVP 演示样本，非正式测绘成果）']:
        p = doc.add_paragraph(text); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def _statistics_table(doc, evidence):
    table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(table.rows[0].cells, ['图层', '要素数', '面积（㎡）', '长度（m）']):
        cell.text = text; _shade(cell, 'E8EEF5'); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for name, item in evidence['layers'].items():
        row = [name, str(item['count']),
               f"{item['area']:.2f}" if item['area'] else '—',
               f"{item['length']:.2f}" if item['length'] else '—']
        for cell, text in zip(table.add_row().cells, row): cell.text = text

def build_report(evidence, sections, ai_status, output_path, template='engineering_summary'):
    """按固定章节结构排版报告：统计数据来自 evidence，叙述段落来自 sections。"""
    if not PYDOCX_OK: raise RuntimeError('未找到 python-docx 运行依赖')
    doc = Document()
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(doc.sections[0], attr, Cm(2.54))
    normal = doc.styles['Normal']; _use_report_font(normal, 10.5)
    normal.paragraph_format.line_spacing = 1.25; normal.paragraph_format.space_after = Pt(5)
    for style_name, size, color in [('Heading 1', 16, '1F4D78'), ('Heading 2', 13, '2E74B5')]:
        _use_report_font(doc.styles[style_name], size, color)
    _cover(doc, template)
    doc.add_heading('一、项目概况', 1)
    doc.add_paragraph(sections['项目概况'])
    doc.add_heading('二、数据来源与处理口径', 1)
    doc.add_paragraph('本报告的数据统计由系统确定性 GIS 计算模块生成。面积按投影平面坐标的多边形鞋带公式计算，'
                      '线长度按相邻顶点欧氏距离累计计算；文本由受控模板生成。')
    doc.add_heading('三、技术路线', 1)
    doc.add_paragraph(sections['技术路线'])
    doc.add_heading('四、成果统计', 1)
    _statistics_table(doc, evidence)
    doc.add_paragraph(f"汇总：共处理 {evidence['feature_count']} 个要素；面状成果面积 {evidence['total_area']:.2f} ㎡"
                      f"（{evidence['total_area']/10000:.4f} ha）；线状成果长度 {evidence['total_length']:.2f} m。", style='Normal')
    doc.add_heading('五、质量检查与评价', 1)
    if evidence['warnings']:
        doc.add_paragraph('系统发现以下需复核项：')
        for item in evidence['warnings']: doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph('基础属性完整性与几何类型检查未发现异常。')
    doc.add_paragraph(sections['质量评价'])
    doc.add_heading('六、结论与建议', 1)
    doc.add_paragraph(sections['结论与建议'])
    doc.add_heading('附录 A：数据血缘记录', 1)
    doc.add_paragraph(f"源文件：用户上传 GIS 矢量成果；空间参考：{evidence['crs']}；生成时间：{date.today()}；"
                      f"处理版本：MVP 0.3；文本生成：{ai_status}。")
    doc.save(output_path)
